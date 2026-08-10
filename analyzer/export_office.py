"""Excel (.xlsx) and Word (.docx) exporters — mirror the Markdown report's
content in Office formats using the same grouped/deduplicated rows."""

import io

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from .report import _group_by_method, _rows_for_method

HEADER_FILL = PatternFill(start_color="17324D", end_color="17324D", fill_type="solid")
HEADER_FONT = Font(color="FFFFFF", bold=True)


def _tipo_label(category: str) -> str:
    if category == "stored_procedure":
        return "Stored Procedure"
    if category == "oracle_package_call":
        return "PL/SQL (Oracle)"
    return "Query"


def _style_header(row_cells):
    for c in row_cells:
        c.font = HEADER_FONT
        c.fill = HEADER_FILL


def build_xlsx(
    app_name, tech, settings, sql_findings, io_findings, security_flags,
    companion_assemblies=None, db_procedures=None, db_tables=None,
) -> bytes:
    wb = Workbook()

    ws = wb.active
    ws.title = "Tecnologia"
    ws.append(["Campo", "Valor"])
    _style_header(ws[1])
    ws.append(["Lenguaje", tech.language])
    ws.append(["Target .NET", tech.dotnet_target])
    ws.append(["UI Framework", ", ".join(tech.ui_framework)])
    ws.append(["Drivers de BD", ", ".join(tech.db_drivers)])
    if companion_assemblies:
        ws.append(["Ensamblados adicionales decompilados", ", ".join(companion_assemblies)])

    def new_sheet(title, headers):
        s = wb.create_sheet(title)
        s.append(headers)
        _style_header(s[1])
        return s

    sec_ws = new_sheet("Seguridad", ["Severidad", "Descripcion", "Ubicacion"])
    for f in security_flags:
        sec_ws.append([f.severity, f.description, f.location])

    conn_ws = new_sheet(
        "Conexiones y config",
        ["Setting", "Valor por defecto", "Archivo", "Categoria", "Extractor", "Linea", "Confianza"],
    )
    for s in settings:
        conn_ws.append([
            s.name, s.default_value, s.source_file, s.category,
            s.evidence.extractor, s.evidence.line_number, s.evidence.confidence,
        ])

    sql_ws = new_sheet(
        "Funciones-SQL-SP",
        ["Clase", "Funcion", "Conexion", "Tipo", "Tabla o SP", "SQL / Query", "Parametros", "Columnas resultado",
         "Extractor", "Linea", "Confianza"],
    )
    for (class_name, method), group in _group_by_method(sql_findings).items():
        for row_text, conn_hint, category, target, params, result_columns, evidence in _rows_for_method(group):
            sql_ws.append([
                class_name, method, conn_hint, _tipo_label(category), target or "?", row_text,
                "\n".join(params) if params else "",
                ", ".join(result_columns) if result_columns else "",
                evidence.extractor, evidence.line_number, evidence.confidence,
            ])
            if params:
                sql_ws.cell(row=sql_ws.max_row, column=7).alignment = Alignment(wrap_text=True, vertical="top")

    io_ws = new_sheet("Archivos-Impresoras-Red", ["Clase", "Funcion", "Operacion", "Detalle"])
    seen = set()
    for f in io_findings:
        key = (f.class_name, f.method, f.raw)
        if key in seen:
            continue
        seen.add(key)
        io_ws.append([f.class_name, f.method, f.operation, f.raw])

    if db_procedures:
        sp_ws = new_sheet("Definiciones SP (BD)", ["Schema.Objeto", "Estado", "Definicion"])
        for p in db_procedures:
            sp_ws.append([
                f"{p['schema_name']}.{p['object_name']}", p["status"], (p.get("definition") or "").strip(),
            ])
            sp_ws.cell(row=sp_ws.max_row, column=3).alignment = Alignment(wrap_text=True, vertical="top")

        param_rows = [
            (f"{p['schema_name']}.{p['object_name']}", prm["name"], prm["type"], prm.get("max_length") or "",
             "Si" if prm["is_output"] else "", prm.get("default") if prm.get("has_default") else "")
            for p in db_procedures for prm in (p.get("parameters") or [])
        ]
        if param_rows:
            params_ws = new_sheet(
                "Parametros SP (BD)", ["Schema.Objeto", "Parametro", "Tipo", "Longitud", "Salida", "Default"],
            )
            for row in param_rows:
                params_ws.append(list(row))

        result_col_rows = [
            (f"{p['schema_name']}.{p['object_name']}", col["name"], col["type"], "Si" if col["nullable"] else "No")
            for p in db_procedures if p.get("result_columns") for col in p["result_columns"]
        ]
        if result_col_rows:
            cols_ws = new_sheet("Columnas resultado SP (BD)", ["Schema.Objeto", "Columna", "Tipo", "Nulo"])
            for row in result_col_rows:
                cols_ws.append(list(row))

    if db_tables:
        tbl_ws = new_sheet("Esquema tablas (BD)", ["Tabla", "Columna", "Tipo", "Nulo", "Longitud", "Default"])
        for t in db_tables:
            full_name = f"{t['schema_name']}.{t['table_name']}"
            for c in t["columns"]:
                tbl_ws.append([
                    full_name, c["name"], c["type"], c["nullable"], c.get("max_length") or "", c.get("default") or "",
                ])
        fk_rows = [
            (f"{t['schema_name']}.{t['table_name']}", fk["column"], fk["ref_table"], fk["ref_column"], fk["fk_name"])
            for t in db_tables for fk in t.get("foreign_keys", [])
        ]
        if fk_rows:
            fk_ws = new_sheet("Claves foraneas (BD)", ["Tabla", "Columna", "Tabla referida", "Columna referida", "FK"])
            for row in fk_rows:
                fk_ws.append(list(row))

    for sheet in wb.worksheets:
        for col_cells in sheet.columns:
            length = max((len(str(c.value)) for c in col_cells if c.value is not None), default=10)
            sheet.column_dimensions[col_cells[0].column_letter].width = min(max(length + 2, 12), 90)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = str(value) if value is not None else ""
    return table


def build_docx(
    app_name, tech, settings, sql_findings, io_findings, security_flags,
    companion_assemblies=None, db_procedures=None, db_tables=None,
) -> bytes:
    doc = Document()
    doc.add_heading(f"{app_name} — Inventario automatico (borrador)", level=1)
    doc.add_paragraph(
        "Generado por QAPV-LegacyAppAnalyzer (ilspycmd + extractor Python). "
        "Este es un primer borrador — revisar antes de usarlo como fuente final."
    )

    doc.add_heading("Tecnologia", level=2)
    tech_rows = [
        ("Lenguaje", tech.language),
        ("Target .NET", tech.dotnet_target),
        ("UI Framework", ", ".join(tech.ui_framework)),
        ("Drivers de BD", ", ".join(tech.db_drivers)),
    ]
    if companion_assemblies:
        tech_rows.append(("Ensamblados adicionales decompilados", ", ".join(companion_assemblies)))
    _add_table(doc, ["Campo", "Valor"], tech_rows)

    if security_flags:
        doc.add_heading("Alertas de seguridad", level=2)
        _add_table(
            doc, ["Severidad", "Descripcion", "Ubicacion"],
            [(f.severity, f.description, f.location) for f in security_flags],
        )

    doc.add_heading("Connection strings y configuraciones", level=2)
    _add_table(
        doc, ["Setting", "Valor por defecto", "Archivo", "Categoria", "Extractor", "Linea", "Confianza"],
        [
            (s.name, s.default_value, s.source_file, s.category,
             s.evidence.extractor, s.evidence.line_number, s.evidence.confidence)
            for s in settings
        ],
    )

    doc.add_heading("Funciones -> SQL / Stored Procedures", level=2)
    sql_rows = []
    for (class_name, method), group in _group_by_method(sql_findings).items():
        for row_text, conn_hint, category, target, params, result_columns, evidence in _rows_for_method(group):
            sql_rows.append((
                class_name, method, conn_hint, _tipo_label(category), target or "?", row_text,
                "; ".join(params) if params else "",
                ", ".join(result_columns) if result_columns else "",
                evidence.extractor, evidence.line_number, evidence.confidence,
            ))
    _add_table(
        doc,
        ["Clase", "Funcion", "Conexion", "Tipo", "Tabla o SP", "SQL / Query", "Parametros", "Columnas resultado",
         "Extractor", "Linea", "Confianza"],
        sql_rows,
    )

    if db_procedures:
        doc.add_heading("Definiciones de Stored Procedures (extraidas de la base de datos, solo lectura)", level=2)
        for p in db_procedures:
            full_name = f"{p['schema_name']}.{p['object_name']}"
            if p["status"] == "ok" and p["definition"]:
                doc.add_heading(full_name, level=3)
                if p.get("parameters"):
                    doc.add_paragraph("Parametros de entrada/salida:")
                    _add_table(
                        doc, ["Parametro", "Tipo", "Longitud", "Salida", "Default"],
                        [
                            (prm["name"], prm["type"], prm.get("max_length") or "",
                             "Si" if prm["is_output"] else "", prm.get("default") if prm.get("has_default") else "")
                            for prm in p["parameters"]
                        ],
                    )
                if p.get("result_columns"):
                    doc.add_paragraph("Columnas que devuelve (determinado por SQL Server, solo lectura):")
                    _add_table(
                        doc, ["Columna", "Tipo", "Nulo"],
                        [(col["name"], col["type"], "Si" if col["nullable"] else "No") for col in p["result_columns"]],
                    )
                elif p.get("result_columns") is None:
                    doc.add_paragraph(
                        "SQL Server no pudo determinar automaticamente las columnas de salida de este SP "
                        "(comun si usa SQL dinamico, tablas temporales o multiples result sets)."
                    )
                para = doc.add_paragraph()
                run = para.add_run((p["definition"] or "").strip())
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            else:
                doc.add_paragraph(f"{full_name} — no disponible ({p['status']})")

    if db_tables:
        doc.add_heading("Esquema de tablas (extraido de la base de datos, solo lectura)", level=2)
        for t in db_tables:
            full_name = f"{t['schema_name']}.{t['table_name']}"
            doc.add_heading(full_name, level=3)
            _add_table(
                doc, ["Columna", "Tipo", "Nulo", "Longitud", "Default"],
                [
                    (c["name"], c["type"], c["nullable"], c.get("max_length") or "", c.get("default") or "")
                    for c in t["columns"]
                ],
            )
            if t.get("foreign_keys"):
                for fk in t["foreign_keys"]:
                    doc.add_paragraph(f"FK: {fk['column']} -> {fk['ref_table']}.{fk['ref_column']} ({fk['fk_name']})")

    doc.add_heading("Archivos, impresoras, procesos y red", level=2)
    io_rows = []
    seen = set()
    for f in io_findings:
        key = (f.class_name, f.method, f.raw)
        if key in seen:
            continue
        seen.add(key)
        io_rows.append((f.class_name, f.method, f.operation, f.raw))
    _add_table(doc, ["Clase", "Funcion", "Operacion", "Detalle"], io_rows)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
